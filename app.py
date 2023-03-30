import camelot
import pandas as pd
import math
import openai
import streamlit as st
import plotly.graph_objs as go
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import base64 

def extract_text_tables(file_path):
    tables = {}
    # Extract tables
    extracted_tables = camelot.read_pdf(file_path, pages='all', flavor='stream')

    for page_number, table in enumerate(extracted_tables):
        # Convert the table to a Pandas DataFrame
        df = table.df

        # Store the table in the dictionary with the page number as the key
        tables[page_number] = df
    return tables

def interpret_job(score): 
    if score >= 130:
        text = "학생은 창조적, 통설적, 전문적인 일에 적합할 수 있으며, 그러한 직업의 예시로는 학자, 교수, 고급공무원 등이 있습니다."
    elif score >=120:
        text = "학생은 지도적, 전문적, 행동적인 일에 적합할 수 있으며, 그러한 직업의 예시로는 의사, 변호사, 작가, 교사 등이 있습니다."    
    elif score >=110:
        text = "학생은 행동적, 지도적인 일에 적합할 수 있으며, 그러한 직업의 예시로는 고급경영자, 고급기술자 등이 있습니다."
    elif score >=90:
        text = "학생은 행동적, 지도적인 일에 적합할 수 있으며, 그러한 직업의 예시로는 각종 상급 기능직 등이 있습니다."
    elif score >=80:
        text = "학생은 행동적, 숙련된 일에 적합할 수 있으며, 그러한 직업의 예시로는 각종 중급 기능직 등이 있습니다."
    elif score >=70:
        text = "학생은 감독하에 행동적으로 수행하는 직종에 적합할 수 있습니다."
    elif score <=69:
        text = "학생은 반복적이고 단순한 일을 수행하는 직종에 적합할 수 있습니다."
    else:
        print("해당되는 범위가 없습니다. 데이터를 다시 확인해보세요.")
    return text

st.set_page_config(
    page_title="보고서", 
    page_icon="📊",
    initial_sidebar_state="expanded"
    )

uploaded_file = st.sidebar.file_uploader('파일 업로드')
uploaded_key = st.sidebar.text_input('OpenAI 키를 입력')

if uploaded_file is not None:

    # byte object into a PDF file 
    with open("input.pdf", "wb") as f:
        base64_pdf = base64.b64encode(uploaded_file.read()).decode('utf-8')
        f.write(base64.b64decode(base64_pdf))
    f.close()

    korean_tables = extract_text_tables("input.pdf")

    df = korean_tables
    df_info = pd.DataFrame(df[0])

    # extract the variables
    examiner = df_info.iloc[0, 0]
    subject = "Test"
    gender = df_info.iloc[1, 1]
    testing_date = df_info.iloc[2, 0]
    age = df_info.iloc[2, 1].split()[0][1:3]
    birth_date = df_info.iloc[2, 1].split()[1][1:-1]
    
    # subscales scores
    subscales_df = pd.DataFrame(df[1])
    subscales_df.columns = ['원점수', '환산점수', '백분위', '추정연령', '측정표준오차(SEM)']
    subscales_df.drop('원점수', axis=1, inplace=True)
    subscales_df.insert(0,'소검사', ['공통성', '어휘', '상식', '이해', '토막짜기', '퍼즐', '행렬추론', '무게비교', '공통그림찾기', '산수', '숫자', '그림기억', '순차연결', '기호쓰기', '동형찾기', '선택'])
    # print(df_subscales)

    # scales scores
    scales_df = pd.DataFrame(df[2])
    ci_per = scales_df.iat[0,3]
    scales_df.drop([0], axis=0, inplace=True)
    scales_df.columns = ['환산점수', '지표점수', '백분위', f'신뢰구간({ci_per}%)', '진단분류(수준)', '측정표준오차(SEM)']
    scales_df.drop('환산점수', axis=1, inplace=True)
    scales_df.insert(0,'지표', ['언어이해', '시공간', '유동추론', '작업기억', '처리속도', '전체IQ'])
    # print(df_scales)

    fsiq_index = scales_df[scales_df["지표"] == "전체IQ"].index[0]
    
    fsiq_score = int(scales_df.loc[fsiq_index, "지표점수"])
    fsiq_percentile = float(scales_df.loc[fsiq_index, "백분위"])
    fsiq_ranking = 100 - math.ceil(fsiq_percentile)
    fsiq_confidence = scales_df.loc[fsiq_index, "신뢰구간(95%)"]
    fsiq_level = scales_df.loc[fsiq_index, "진단분류(수준)"]
    body = f'''K-WISC-V(한국 웩슬러 아동 지능검사 5판)로 추정한 현재 지능은 '{fsiq_level}'(FSIQ: {fsiq_score}) 범위에 속하는 하며, 백분위가 {math.ceil(fsiq_percentile)}%ile로 나타났다. 이는 {fsiq_score}점 아래에 있는 학생들이 전체 중 {math.ceil(fsiq_percentile)}%가 있다는 것을 의미합니다. 좀 더 쉽게 설명하면, 전체 100명 중 {math.ceil(fsiq_percentile)}번째에 위치 하는 것이며, 앞에서 {math.ceil(fsiq_ranking)}등이라고 할 수 있습니다.
    95% 신뢰구간은 {fsiq_confidence}으로 학생이 검사를 100번 실시했을 때 95번은 {fsiq_confidence}사이의 점수를 받게 된다는 의미입니다. 학생의 컨디션이나 기타 환경적 요소 때문에 어떤 날은 좀 더 높은 점수를 받을 수 있고, 어떤 날은 좀 더 낮은 점수를 받을 수 있습니다. 대부분의 경우는 두 점수 사이의 평균적인 점수를 받게 될 것입니다. '''

    #!----- Plotly -----
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=scales_df['지표'],
        y=scales_df['지표점수'],
        mode='lines+markers+text',
        name='K-WISC-V',
        line=dict(color='green'),
        text=scales_df['지표점수'],
        textposition="top center"
    ))
    fig.update_yaxes(title_text=None, showgrid=False, showline=True, linecolor='lightgrey', linewidth=2)
    fig.update_xaxes(title_text=None, showgrid=False, showline=True, linecolor='lightgrey', linewidth=2)
    fig.update_layout(
        plot_bgcolor='rgba(0,0,0,0)',
        yaxis_range=[40, 160],
        margin=dict(l=0, r=0, t=50, b=40)
    )
    

    #!----- Streamlit -----
    st.header('연구 참가자 K-WISC-V 보고서')
    st.markdown('')
    col1, col2 = st.columns(2)
    with col1:
        st.markdown(f'> 인적사항: {subject} ({gender}, {age})')
    with col2:
        st.markdown(f'> 검사일: {testing_date}')
    
    st.markdown('')
    st.markdown('')
    
    st.markdown('##### 지표 점수')
    st.dataframe(scales_df, use_container_width=True)
    
    with st.expander('지표 점수 차트'):
        st.plotly_chart(fig, use_container_width=True)
    
    st.markdown('')
    
    st.markdown('##### 소검사 점수')
    st.dataframe(subscales_df, use_container_width=True)
    
    st.markdown('')
    
    st.markdown('##### 해석')
    interpretation = st.text_area("해석",
                 body+interpret_job(fsiq_score),
                 height=200, 
                 label_visibility="collapsed")
    generate = st.checkbox("보고서 생성")
    if generate == True:
        # Load the report.docx file using python-docx
        document = Document('report.docx')
        # Replace the markers with the corresponding values
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '<<name>>' in cell.text:
                        cell.text = cell.text.replace('<<name>>', subject)
                    if '<<age>>' in cell.text:
                        cell.text = cell.text.replace('<<age>>', age)
                    if '<<test_date>>' in cell.text:
                        cell.text = cell.text.replace('<<test_date>>', testing_date)
                    if '<<gender>>' in cell.text:
                        cell.text = cell.text.replace('<<gender>>', gender)
                    if '<<grade>>' in cell.text:
                        cell.text = cell.text.replace('<<grade>>', '')
                    if '<<birth_date>>' in cell.text:
                        cell.text = cell.text.replace('<<birth_date>>', birth_date)
                    if '<<picture>>' in cell.text:
                        # Save the plotly chart to a temporary file
                        with tempfile.NamedTemporaryFile(suffix='.png') as tmp_file:
                            fig.write_image(tmp_file.name, scale=2)
                            # Insert the plotly chart as a picture
                            cell.text = cell.text.replace('<<picture>>', '')
                            run = cell.paragraphs[0].add_run()
                            run.add_picture(tmp_file.name, width=Inches(6), height=Inches(3))
                    if '<<interpretation>>' in cell.text:
                        # Insert the interpretation text
                        cell.text = cell.text.replace('<<interpretation>>', '')
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        paragraph.add_run(interpretation)

        # Offer the report.docx file for download with the streamlit download button
        document.save('modified_report.docx')

        st.download_button(
            label="보고서 받기",
            data=open('modified_report.docx', 'rb').read(),
            file_name=f"{subject} 연구 참가자 보고서.docx",
            mime="application/octet-stream"
        )

    #!----- OpenAI -----
    openai.api_key = uploaded_key
    prompt = f'''
    {scales_df}
    {subscales_df}
    
    작성방법
    - 위 데이터는 WISC-V검사의 결과이며, 전체 IQ는 무조건 처음 언급해준다. 작성 순서는 전체 IQ, 다른 지표 점수, 소검사 순으로 해석하고 작성한다.
    - 실시를 하지 않은 검사는 언급할 필요가 없다.
    - 소검사의 환산점수의 평균은 10점이다.
    '''
    if st.button("AI 분석", type="primary"):
        st.markdown("----")
        res_box = st.empty()
        report = []
        # Looping over the response
        output = openai.ChatCompletion.create(model='gpt-3.5-turbo',
                                            messages=[{"role": "system", "content": "너는 임상 심리 전문가이다. 심리검사 점수는 정확하게 참고해야한다."},
                                                      {"role": "user", "content":prompt}],
                                            max_tokens=2000, 
                                            temperature = 0.4)
        for item in output['choices']:
            chatgpt_output = item['message']['content']
        st.markdown(f''' ##### Chat-GPT 답변
{chatgpt_output}''')
        st.markdown("----")
        
